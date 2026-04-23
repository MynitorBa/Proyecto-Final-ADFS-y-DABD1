#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Elimina secciones SQL de los 3 manuales Word (Parte 5 + Anexo métricas).
Detección de límites de sección basada en TEXTO, no en tamaño de fuente,
para evitar falsos negativos por herencia de estilos.
"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import os

BASE = r"C:\Proyecto-Final-ADFS-y-DABD1"

# ─── helpers ─────────────────────────────────────────────────────────────────

def is_page_break_para(para):
    return 'w:type="page"' in para._element.xml

def txt(para):
    return para.text.strip()

# Patrones de texto que identifican CUALQUIER sección h1 del documento.
# Usamos texto, no tamaño de fuente, para máxima robustez.
H1_PREFIXES = [
    'parte 0', 'parte 1', 'parte 2', 'parte 3', 'parte 4',
    'parte 5', 'parte 6', 'parte 7', 'parte 8', 'parte 9',
    'anexo', 'inventario de cobertura',
    'universidad del istmo', 'manual del programador',
    'índice', 'indice',
]

def is_any_section_boundary(text):
    """True si el texto empieza con un patrón conocido de sección h1."""
    t = text.lower().strip()
    return any(t.startswith(pfx) for pfx in H1_PREFIXES)

# Patrones de CONTENIDO que identifican secciones SQL a eliminar.
# Deliberadamente NO incluimos 'parte 5' como prefijo genérico,
# para no confundir con 'Parte 5 — Sistema de Auditoría' (Agencia).
SQL_CONTENT_MARKERS = [
    'queries sql',
    'queries de m\u00e9tricas',   # "queries de métricas"
    'queries de metricas',
    'm\u00e9tricas para defensa',  # "métricas para defensa"
    'metricas para defensa',
]

def is_sql_section(text):
    """True si el límite de sección corresponde a una sección SQL a eliminar."""
    t = text.lower()
    # Anexo: siempre eliminamos
    if t.startswith('anexo'):
        return True
    # Secciones cuyo CONTENIDO es SQL (por nombre)
    if any(marker in t for marker in SQL_CONTENT_MARKERS):
        return True
    return False

def apply_renames_to_runs(para, renames):
    for run in para.runs:
        for old, new in renames.items():
            if old in run.text:
                run.text = run.text.replace(old, new)

def apply_renames_to_tables(doc, renames):
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    apply_renames_to_runs(para, renames)

# ─── core ────────────────────────────────────────────────────────────────────

def process(fname, renames=None):
    path = os.path.join(BASE, fname)
    doc  = Document(path)
    paras = list(doc.paragraphs)

    to_delete = set()
    in_sql    = False

    for i, p in enumerate(paras):
        t = txt(p)

        if is_any_section_boundary(t):
            if is_sql_section(t):
                # Sección SQL: marcar para borrar + page_break previo
                in_sql = True
                to_delete.add(i)
                if i > 0 and is_page_break_para(paras[i - 1]):
                    to_delete.add(i - 1)
            else:
                # Sección NO-SQL: detener borrado, aplicar renombres
                in_sql = False
                if renames:
                    apply_renames_to_runs(p, renames)
        elif in_sql:
            # Dentro de sección SQL: borrar todo
            to_delete.add(i)
        else:
            # Párrafo normal fuera de sección SQL: sólo aplicar renombres
            if renames:
                apply_renames_to_runs(p, renames)

    # Aplicar renombres en tablas (para Inventario de Cobertura, etc.)
    if renames:
        apply_renames_to_tables(doc, renames)

    # Eliminar en orden inverso para no corromper índices
    for i in sorted(to_delete, reverse=True):
        el = paras[i]._element
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)

    doc.save(path)

    # Recolectar secciones h1 restantes (por texto, consistente con detección)
    remaining = []
    for p in doc.paragraphs:
        t = txt(p)
        if is_any_section_boundary(t) and t:
            # excluir portada y meta
            if not any(t.lower().startswith(x) for x in
                       ['universidad del istmo', 'manual del programador',
                        'índice', 'indice']):
                remaining.append(t)
    return len(to_delete), remaining

# ─── main ────────────────────────────────────────────────────────────────────

DOCS = {
    "Manual_Programador_Aerolinea.docx": None,
    "Manual_Programador_Hotelera.docx":  None,
    # Para Agencia: Parte 6 (Auditoría) → Parte 5; subsecciones 6.x → 5.x
    "Manual_Programador_Agencia.docx": {
        "Parte 6 \u2014": "Parte 5 \u2014",   # em dash
        "Parte 6 -":     "Parte 5 -",
        "6.1 ":          "5.1 ",
        "6.2 ":          "5.2 ",
        "6.3 ":          "5.3 ",
        "Parte 6,":      "Parte 5,",
        "Parte 6 ":      "Parte 5 ",
    },
}

for fname, renames in DOCS.items():
    n, sections = process(fname, renames)
    label = fname.replace("Manual_Programador_", "").replace(".docx", "")
    print(f"\n{'='*60}")
    print(f"  {label}")
    print(f"{'='*60}")
    print(f"  Parrafos eliminados : {n}")
    print(f"  Secciones h1 restantes:")
    for s in sections:
        try:
            print(f"    * {s}")
        except UnicodeEncodeError:
            print(f"    * (titulo con caracteres especiales)")

print("\n[OK] Los 3 documentos fueron modificados y guardados.")
