#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Utilidades comunes para generación de documentos Word."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

COLOR_H1   = RGBColor(0x1F, 0x38, 0x64)
COLOR_H2   = RGBColor(0x2E, 0x75, 0xB6)
COLOR_WHITE= RGBColor(0xFF, 0xFF, 0xFF)
HEX_CODE_BG      = 'F2F2F2'
HEX_TABLE_HEADER = '2E75B6'
HEX_TABLE_ALT    = 'D9E2F3'

def new_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width   = Inches(8.5)
    sec.page_height  = Inches(11)
    sec.left_margin  = Inches(1)
    sec.right_margin = Inches(1)
    sec.top_margin   = Inches(1)
    sec.bottom_margin= Inches(1.2)
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    return doc

def set_footer(doc, text="Universidad del Istmo (UNIS) — Documentación Técnica"):
    sec = doc.sections[0]
    footer = sec.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)

def add_cover(doc, module_name):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Universidad del Istmo (UNIS)")
    r.font.name = 'Calibri'; r.font.size = Pt(16); r.font.bold = True
    r.font.color.rgb = COLOR_H1

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Documentación Técnica  |  Proyecto Final ADFS y DABD1")
    r.font.name = 'Calibri'; r.font.size = Pt(12)
    r.font.color.rgb = COLOR_H2

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Manual del Programador\n{module_name}")
    r.font.name = 'Calibri'; r.font.size = Pt(26); r.font.bold = True
    r.font.color.rgb = COLOR_H1

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Guía técnica interna para equipo de desarrollo")
    r.font.name = 'Calibri'; r.font.size = Pt(13); r.font.italic = True

    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("19 de abril de 2026")
    r.font.name = 'Calibri'; r.font.size = Pt(12)

    doc.add_page_break()

def h1(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Calibri'; r.font.size = Pt(16); r.font.bold = True
    r.font.color.rgb = COLOR_H1
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    return p

def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Calibri'; r.font.size = Pt(13); r.font.bold = True
    r.font.color.rgb = COLOR_H2
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    return p

def h3(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Calibri'; r.font.size = Pt(11); r.font.bold = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    return p

def body(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Calibri'; r.font.size = Pt(11)
    return p

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.font.bold = True; rb.font.name = 'Calibri'; rb.font.size = Pt(11)
    r = p.add_run(text)
    r.font.name = 'Calibri'; r.font.size = Pt(11)
    return p

def code(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Consolas'; r.font.size = Pt(9)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), HEX_CODE_BG)
    pPr.append(shd)
    p.paragraph_format.left_indent  = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    return p

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    # Header
    hr = t.rows[0]
    for i, h in enumerate(headers):
        c = hr.cells[i]
        c.text = h
        if c.paragraphs[0].runs:
            rn = c.paragraphs[0].runs[0]
            rn.font.name = 'Calibri'; rn.font.size = Pt(10)
            rn.font.bold = True; rn.font.color.rgb = COLOR_WHITE
        set_cell_bg(c, HEX_TABLE_HEADER)
    # Rows
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        bg = HEX_TABLE_ALT if ri % 2 == 1 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            c = row.cells[ci]
            c.text = str(val) if val is not None else ''
            if c.paragraphs[0].runs:
                rn = c.paragraphs[0].runs[0]
                rn.font.name = 'Calibri'; rn.font.size = Pt(10)
            set_cell_bg(c, bg)
    return t

def page_break(doc):
    doc.add_page_break()
